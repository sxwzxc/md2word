"""
Python Cloud Function - Markdown to DOCX Converter
api/convert/md-to-docx.py → POST /api/convert/md-to-docx
Accepts a markdown text body and returns the .docx file as a base64-encoded
JSON response (so the binary survives any transport re-encoding on the platform).

Query param ?mode=1 (default, basic) or ?mode=2 (elegant typography).
"""
import re
import json
import base64
from io import BytesIO
from http.server import BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r'(\*\*[^*]+\*\*'          # **bold**
    r'|\*\([^*]+\)\*'          # *(italic)*  (rare, keep simple)
    r'|\*[^*\s][^*]*\*'        # *italic*
    r'|`[^`]+`'                # `inline code`
    r'|\[[^\]]+\]\([^)]+\)'    # [text](url)
    r'|~~[^~]+~~)'             # ~~strikethrough~~
)


# East Asian font name — ensures Chinese / Japanese / Korean characters
# render with a real font instead of showing as missing-glyph boxes (叉叉).
_EA_FONT = '微软雅黑'

# Heading East-Asian font — 黑体 (SimHei) is the standard Chinese
# paper / report heading typeface.
_HEADING_EA_FONT = '黑体'

# Soft body text color for elegant mode (instead of harsh pure black).
_BODY_COLOR = RGBColor(0x33, 0x33, 0x33)


def _strip_theme_fonts(r_fonts):
    """Remove *Theme font attributes so explicit font names take precedence.

    The python-docx default template defines styles and docDefaults with
    asciiTheme/eastAsiaTheme/hAnsiTheme that point at theme fonts whose
    East-Asian slot is EMPTY.  Word/WPS then cannot find a CJK font and
    renders missing-glyph squares (方框点) in front of text.  Stripping
    these attributes lets our explicit ascii=Calibri / eastAsia=微软雅黑 win.
    """
    for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        if r_fonts.get(qn(attr)) is not None:
            del r_fonts.attrib[qn(attr)]


def _set_run_fonts(run, ascii_font='Calibri', ea_font=_EA_FONT):
    """Set both ASCII and East-Asian fonts on a run so CJK text renders."""
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    _strip_theme_fonts(r_fonts)
    r_fonts.set(qn('w:eastAsia'), ea_font)
    r_fonts.set(qn('w:ascii'), ascii_font)
    r_fonts.set(qn('w:hAnsi'), ascii_font)


def _add_shading(paragraph, fill="F2F2F2"):
    """Apply background shading to a paragraph (used for code blocks)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def _add_paragraph_border(paragraph, side='left', color='3776AB', sz='24'):
    """Add a single-side colored border to a paragraph (for blockquotes)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), sz)       # eighths of a point; 24 = 3pt
    bdr.set(qn('w:space'), '8')   # gap from text (twips)
    bdr.set(qn('w:color'), color)
    p_bdr.append(bdr)


def _set_paragraph_spacing(paragraph, before=None, after=None, line=None):
    """Set paragraph spacing. before/after in Pt, line as a multiple (e.g. 1.5)."""
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _add_hyperlink(paragraph, text, url, color='0563C1'):
    """Add a real, clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:eastAsia'), _EA_FONT)
    r_fonts.set(qn('w:ascii'), 'Calibri')
    r_fonts.set(qn('w:hAnsi'), 'Calibri')
    r_pr.append(r_fonts)

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    r_pr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)

    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_formatted_runs(paragraph, text, doc=None, mode=1):
    """Parse inline markdown and add styled runs to *paragraph*."""
    if not text:
        return
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        # Bold
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            _set_run_fonts(run)
            if mode == 2:
                run.font.color.rgb = _BODY_COLOR
        # Strikethrough
        elif part.startswith('~~') and part.endswith('~~') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
            _set_run_fonts(run)
        # Inline code
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            _set_run_fonts(run, ascii_font='Consolas')
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        # Link [text](url)
        elif part.startswith('[') and ']' in part and '(' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                link_text, url = m.group(1), m.group(2)
                if mode == 2 and doc is not None:
                    _add_hyperlink(paragraph, link_text, url)
                else:
                    run = paragraph.add_run(link_text)
                    _set_run_fonts(run)
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                    run.underline = True
            else:
                run = paragraph.add_run(part)
                _set_run_fonts(run)
        # Italic (single *)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            _set_run_fonts(run)
        else:
            run = paragraph.add_run(part)
            _set_run_fonts(run)
            if mode == 2:
                run.font.color.rgb = _BODY_COLOR


# ---------------------------------------------------------------------------
# Block-level markdown → docx conversion
# ---------------------------------------------------------------------------

def _shade_cell(cell, fill):
    """Apply background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def _convert_table(doc, header_row, data_rows, mode=1):
    """Add a markdown table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    # Header
    for idx, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[idx]
        cell.text = ''
        if mode == 2:
            _shade_cell(cell, '2B579A')  # Word blue header
        p = cell.paragraphs[0]
        if mode == 2:
            p.alignment = 1  # center
        _add_formatted_runs(p, cell_text.strip(), doc=doc, mode=mode)
        for run in p.runs:
            run.bold = True
            _set_run_fonts(run)
            if mode == 2:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data
    for row_idx, row_cells in enumerate(data_rows):
        cells = table.add_row().cells
        for idx in range(cols):
            cell = cells[idx]
            cell.text = ''
            if mode == 2:
                # Zebra striping
                _shade_cell(cell, 'F2F6FC' if row_idx % 2 == 0 else 'FFFFFF')
            _add_formatted_runs(
                cell.paragraphs[0],
                row_cells[idx].strip() if idx < len(row_cells) else '',
                doc=doc, mode=mode,
            )


def _purge_heading_styles(doc):
    """Delete ALL heading / title / subtitle / TOC / list styles from
    styles.xml.

    The python-docx default template ships Heading1-9, Title, Subtitle,
    TOCHeading, ListBullet1-3, ListNumber1-3 and their linked *Char
    character styles.  Even when no paragraph references them, WPS Office
    picks up the ``<w:numPr>`` inside the List styles and renders broken
    bullet glyphs (empty ``<w:lvlText>`` + Symbol font) as SQUARE BOXES
    in front of text.  Removing every one of these style definitions
    guarantees a completely clean document with no list leakage.

    We use manual text prefixes (``•`` / ``1.``) for lists, so no list
    style or numbering definition is needed at all.
    """
    styles_element = doc.styles.element
    to_remove = []
    for style in styles_element.findall(qn('w:style')):
        style_id = style.get(qn('w:styleId')) or ''
        name_el = style.find(qn('w:name'))
        name_val = name_el.get(qn('w:val')) if name_el is not None else ''
        id_lower = style_id.lower()
        name_lower = name_val.lower()
        keywords = ('heading', 'title', 'subtitle', 'toc', 'list')
        if any(kw in id_lower or kw in name_lower for kw in keywords):
            to_remove.append(style)
    for style in to_remove:
        styles_element.remove(style)


def _strip_all_numbering(doc):
    """Remove ``<w:numPr>`` from EVERY style and document defaults —
    eliminates any inherited list/bullet formatting anywhere.
    """
    styles_element = doc.styles.element
    for style in styles_element.findall(qn('w:style')):
        p_pr = style.find(qn('w:pPr'))
        if p_pr is not None:
            for num_pr in p_pr.findall(qn('w:numPr')):
                p_pr.remove(num_pr)
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        p_pr_default = doc_defaults.find(qn('w:pPrDefault'))
        if p_pr_default is not None:
            p_pr = p_pr_default.find(qn('w:pPr'))
            if p_pr is not None:
                for num_pr in p_pr.findall(qn('w:numPr')):
                    p_pr.remove(num_pr)


def _sanitize_all_fonts(doc):
    """Strip *Theme font references from EVERY style and docDefaults, and
    set explicit eastAsia/ascii/hAnsi fonts.

    This is the critical fix for the "方框点" (square dot) issue: the
    default theme1.xml has an EMPTY East-Asian font slot
    (``<a:ea typeface=""/>``), so any style referencing
    ``eastAsiaTheme="minorEastAsia"`` ends up with no CJK font and
    renders missing-glyph squares.  By removing every *Theme attribute
    and writing explicit font names, we guarantee CJK text always renders.
    """
    styles_element = doc.styles.element

    def _fix_rfonts(r_fonts):
        if r_fonts is None:
            return
        _strip_theme_fonts(r_fonts)
        # Ensure explicit east-Asian font is set if not already present.
        if r_fonts.get(qn('w:eastAsia')) is None:
            r_fonts.set(qn('w:eastAsia'), _EA_FONT)
        if r_fonts.get(qn('w:ascii')) is None:
            r_fonts.set(qn('w:ascii'), 'Calibri')
        if r_fonts.get(qn('w:hAnsi')) is None:
            r_fonts.set(qn('w:hAnsi'), 'Calibri')

    # 1) docDefaults / rPrDefault
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        r_pr_default = doc_defaults.find(qn('w:rPrDefault'))
        if r_pr_default is not None:
            r_pr = r_pr_default.find(qn('w:rPr'))
            if r_pr is not None:
                _fix_rfonts(r_pr.find(qn('w:rFonts')))

    # 2) Every rFonts element ANYWHERE in styles.xml — this catches rFonts
    #    inside <w:style><w:rPr>, inside <w:tblStylePr><w:rPr>, inside
    #    linked <w:link><w:rPr>, etc.  Using iter() (recursive) not find().
    for r_fonts in styles_element.iter(qn('w:rFonts')):
        _fix_rfonts(r_fonts)


def _patch_theme_fonts(doc):
    """Patch word/theme/theme1.xml so its empty East-Asian font slots
    point to a real CJK font.  This is a belt-and-suspenders fix: even if
    some renderer ignores our style-level font overrides and falls back
    to the theme, it will now find a valid CJK font instead of an empty
    string (which produces the square missing-glyph dots).
    """
    # The theme part is related from the DOCUMENT part (not the package).
    theme_reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme'
    theme_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == theme_reltype:
            theme_part = rel.target_part
            break
    if theme_part is None:
        return
    try:
        blob = theme_part.blob
    except Exception:
        return
    if not blob:
        return
    try:
        xml = blob.decode('utf-8')
    except Exception:
        return
    # Replace empty <a:ea typeface=""/> with 微软雅黑.
    new_xml = re.sub(
        r'(<a:ea\s+typeface=")\s*(")',
        rf'\1{_EA_FONT}\2',
        xml,
    )
    # Also fill empty <a:cs typeface=""/> with a safe default.
    new_xml = re.sub(
        r'(<a:cs\s+typeface=")\s*(")',
        r'\1Calibri\2',
        new_xml,
    )
    if new_xml != xml:
        new_blob = new_xml.encode('utf-8')
        # python-docx Part stores the binary payload in _blob; the blob
        # property is a read-only accessor in some versions.  Set _blob
        # directly so doc.save() picks up the patched content.
        try:
            theme_part._blob = new_blob
        except Exception:
            try:
                theme_part.blob = new_blob
            except Exception:
                pass


def _purge_numbering_part(doc):
    """Empty word/numbering.xml so NO bullet / number definitions exist.

    The default template's numbering.xml contains ``<w:abstractNum>``
    entries whose bullet levels use an EMPTY ``<w:lvlText w:val=""/>``
    with ``<w:rFonts w:ascii="Symbol">``.  WPS Office renders these
    broken bullet definitions as SQUARE BOXES (方框点) in front of text.
    Even after deleting the List styles that reference them, WPS can
    still pick up the numbering definitions directly.  Removing every
    ``<w:abstractNum>`` and ``<w:num>`` child from the numbering
    element eliminates every bullet / number definition so there is
    nothing left to render.

    We use manual text prefixes (``•`` / ``1.``) for lists, so no
    numbering definition is needed at all.

    NOTE: NumberingPart is an XmlPart whose ``blob`` property
    serialises from ``_element`` — setting ``_blob`` has no effect.
    We must mutate the lxml element tree directly.
    """
    numbering_reltype = (
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
    )
    numbering_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == numbering_reltype:
            numbering_part = rel.target_part
            break
    if numbering_part is None:
        return
    try:
        numbering_el = numbering_part.element
    except Exception:
        return
    # Remove ALL children (abstractNum, num, etc.) so the part becomes
    # an empty <w:numbering/> wrapper.
    for child in list(numbering_el):
        numbering_el.remove(child)


def _setup_styles(doc, mode=1):
    """Configure ONLY the Normal (body) style.

    Heading / Title styles are deliberately NOT configured here — they are
    deleted entirely by ``_purge_heading_styles`` so they can never leak
    dots or bullets into the document.
    """
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    if mode == 2:
        normal.font.color.rgb = _BODY_COLOR
        pf = normal.paragraph_format
        pf.line_spacing = 1.5
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.space_after = Pt(6)
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    _strip_theme_fonts(r_fonts)
    r_fonts.set(qn('w:eastAsia'), _EA_FONT)
    r_fonts.set(qn('w:ascii'), 'Calibri')
    r_fonts.set(qn('w:hAnsi'), 'Calibri')


def _setup_page(doc, mode=1):
    """Set page margins. Mode 2 uses comfortable 2.54 cm margins."""
    if mode != 2:
        return
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def _add_heading(doc, level, text, mode=1):
    """Add a heading as a PLAIN paragraph with explicit run formatting.

    No heading / title style is used anywhere — not as a paragraph style
    reference (pStyle), not as a style definition in styles.xml, and not
    via outlineLvl.  The paragraph is completely clean: only spacing
    properties in pPr, and bold + size + color + font on every run.

    NOTE: ``keep_with_next`` and ``keep_together`` are deliberately NOT
    set.  In WPS Office, ``<w:keepNext/>`` is displayed as a SMALL SQUARE
    mark (▪) at the left margin of the paragraph when formatting marks
    are visible, and "Clear Format" removes it.  This was the root cause
    of the persistent "方形点" (square dot) before headings.

    Typography follows the common Chinese paper / report convention:
    黑体 (SimHei) bold, with sizes close to 二号/三号/小三/四号/小四/五号.
    """
    p = doc.add_paragraph()

    # Heading sizes / colors / spacing per level.
    if mode == 2:
        heading_cfg = {
            1: (Pt(22), RGBColor(0x1A, 0x2A, 0x4F), 18, 8),  # 二号  深蓝
            2: (Pt(18), RGBColor(0x2B, 0x57, 0x9A), 14, 6),  # 三号  中蓝
            3: (Pt(15), RGBColor(0x37, 0x76, 0xAB), 10, 4),  # 小三  浅蓝
            4: (Pt(13), RGBColor(0x44, 0x55, 0x66), 8, 4),   # 四号  深灰蓝
            5: (Pt(12), RGBColor(0x55, 0x55, 0x55), 6, 3),   # 小四  中灰
            6: (Pt(11), RGBColor(0x66, 0x66, 0x66), 6, 3),   # 五号  浅灰
        }
    else:
        heading_cfg = {
            1: (Pt(22), RGBColor(0x1F, 0x38, 0x64), 12, 6),
            2: (Pt(18), RGBColor(0x2E, 0x4B, 0x8B), 10, 5),
            3: (Pt(15), RGBColor(0x37, 0x76, 0xAB), 8, 4),
            4: (Pt(13), RGBColor(0x44, 0x44, 0x44), 6, 3),
            5: (Pt(12), RGBColor(0x66, 0x66, 0x66), 6, 3),
            6: (Pt(11), RGBColor(0x66, 0x66, 0x66), 6, 3),
        }
    size, color, before, after = heading_cfg.get(level, heading_cfg[6])

    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

    # Build the heading text runs (preserves inline **bold**, `code`, links).
    _add_formatted_runs(p, text, doc=doc, mode=mode)

    # Override EVERY run with explicit heading formatting:
    #   bold (w:b + w:bCs), size (w:sz + w:szCs), color, and 黑体 font.
    for run in p.runs:
        run.bold = True
        run.font.size = size
        run.font.color.rgb = color
        r_pr = run._element.get_or_add_rPr()
        # Ensure bCs (bold for complex scripts) is set alongside w:b.
        if r_pr.find(qn('w:bCs')) is None:
            r_pr.append(OxmlElement('w:bCs'))
        # szCs (complex-script size) matching w:sz.
        sz_cs = r_pr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            r_pr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(int(size.pt * 2)))
        # Fonts: 黑体 for East-Asian, Calibri for Latin.
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = OxmlElement('w:rFonts')
            r_pr.append(r_fonts)
        _strip_theme_fonts(r_fonts)
        r_fonts.set(qn('w:eastAsia'), _HEADING_EA_FONT)
        r_fonts.set(qn('w:ascii'), 'Calibri')
        r_fonts.set(qn('w:hAnsi'), 'Calibri')

    return p


def convert_markdown_to_docx(markdown_text, mode=1):
    """Convert markdown text to a .docx byte string."""
    doc = Document()

    # Configure Normal style only (no heading styles).
    _setup_styles(doc, mode=mode)
    _setup_page(doc, mode=mode)
    # Delete ALL heading / title / subtitle / TOC / list style definitions
    # so they can never leak dots or bullets into the document.
    _purge_heading_styles(doc)
    # Strip numPr from EVERY style and from docDefaults.
    _strip_all_numbering(doc)
    # Empty numbering.xml so NO bullet / number definitions exist at all.
    _purge_numbering_part(doc)
    # Strip ALL *Theme font references from styles + docDefaults and fill
    # in explicit eastAsia/ascii/hAnsi fonts.  This eliminates the empty
    # eastAsiaTheme that produces missing-glyph square dots.
    _sanitize_all_fonts(doc)
    # Patch theme1.xml so its empty <a:ea typeface=""/> slots point to a
    # real CJK font — belt-and-suspenders against renderers that ignore
    # style-level overrides and fall straight back to the theme.
    _patch_theme_fonts(doc)

    lines = markdown_text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
    i = 0
    n = len(lines)
    first_heading = None
    list_counter = 0  # for ordered lists

    while i < n:
        line = lines[i]

        # ---- Fenced code block ----
        stripped = line.strip()
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            if mode == 2:
                _add_shading(p, 'F5F5F5')
                _add_paragraph_border(p, 'left', '3776AB', '18')
                _set_paragraph_spacing(p, before=6, after=6, line=1.0)
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                _add_shading(p)
            run = p.add_run('\n'.join(code_lines))
            _set_run_fonts(run, ascii_font='Consolas')
            run.font.size = Pt(9.5)
            continue

        # ---- Heading ----
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            if first_heading is None and level == 1:
                first_heading = heading_text
            _add_heading(doc, level, heading_text, mode=mode)
            i += 1
            continue

        # ---- Horizontal rule ----
        if re.match(r'^\s*([-*_])\1{2,}\s*$', line):
            p = doc.add_paragraph()
            if mode == 2:
                _add_paragraph_border(p, 'bottom', 'CCCCCC', '6')
                _set_paragraph_spacing(p, before=8, after=8)
            else:
                _add_shading(p, 'CCCCCC')
            i += 1
            continue

        # ---- Table ----
        if '|' in line and i + 1 < n and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i + 1]):
            header_cells = [c for c in line.strip().strip('|').split('|')]
            i += 2  # skip separator
            data_rows = []
            while i < n and '|' in lines[i] and lines[i].strip():
                data_rows.append([c for c in lines[i].strip().strip('|').split('|')])
                i += 1
            _convert_table(doc, header_cells, data_rows, mode=mode)
            continue

        # ---- Blockquote ----
        if stripped.startswith('>'):
            quote_lines = []
            while i < n and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^\s*>\s?', '', lines[i]))
                i += 1
            quote_text = '\n'.join(quote_lines)
            p = doc.add_paragraph()
            if mode == 2:
                _add_paragraph_border(p, 'left', '3776AB', '24')
                _set_paragraph_spacing(p, before=6, after=6, line=1.4)
                p.paragraph_format.left_indent = Cm(0.6)
                run = p.add_run(quote_text)
                _set_run_fonts(run)
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                run = p.add_run(quote_text)
                _set_run_fonts(run)
            continue

        # ---- Unordered list (manual bullet prefix — no ListBullet style) ----
        if re.match(r'^\s*[-*+]\s+', line):
            while i < n and re.match(r'^\s*[-*+]\s+', lines[i]):
                text = re.sub(r'^\s*[-*+]\s+', '', lines[i])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75)
                p.paragraph_format.first_line_indent = Cm(-0.4)
                _add_formatted_runs(p, '• ' + text, doc=doc, mode=mode)
                if mode == 2:
                    _set_paragraph_spacing(p, before=2, after=2, line=1.4)
                i += 1
            continue

        # ---- Ordered list (manual number prefix — no ListNumber style) ----
        if re.match(r'^\s*\d+\.\s+', line):
            list_counter = 0
            while i < n and re.match(r'^\s*\d+\.\s+', lines[i]):
                list_counter += 1
                text = re.sub(r'^\s*\d+\.\s+', '', lines[i])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75)
                p.paragraph_format.first_line_indent = Cm(-0.4)
                _add_formatted_runs(p, f'{list_counter}. ' + text, doc=doc, mode=mode)
                if mode == 2:
                    _set_paragraph_spacing(p, before=2, after=2, line=1.4)
                i += 1
            continue

        # ---- Blank line ----
        if not stripped:
            i += 1
            continue

        # ---- Regular paragraph (gather consecutive non-empty, non-special lines) ----
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if (not nxt.strip()
                    or nxt.strip().startswith('```')
                    or re.match(r'^#{1,6}\s', nxt)
                    or re.match(r'^\s*[-*+]\s+', nxt)
                    or re.match(r'^\s*\d+\.\s+', nxt)
                    or nxt.strip().startswith('>')
                    or re.match(r'^\s*([-*_])\1{2,}\s*$', nxt)):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _add_formatted_runs(p, ' '.join(l.strip() for l in para_lines), doc=doc, mode=mode)

    # Set document title from first H1 (metadata only — not a visible element)
    if first_heading:
        try:
            doc.core_properties.title = first_heading
        except Exception:
            pass

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTTP handler
# ---------------------------------------------------------------------------

class handler(BaseHTTPRequestHandler):
    """POST /api/convert/md-to-docx — convert markdown body to base64-encoded .docx JSON."""

    def do_POST(self):
        try:
            parsed = urlparse(self.path)
            qs = parse_qs(parsed.query)
            mode_str = qs.get('mode', ['1'])[0]
            mode = 2 if mode_str == '2' else 1

            length = int(self.headers.get('Content-Length', 0))
            raw = self.rfile.read(length) if length else b''
            try:
                markdown_text = raw.decode('utf-8')
            except UnicodeDecodeError:
                markdown_text = raw.decode('utf-8', errors='replace')

            if not markdown_text.strip():
                self._json_error(400, "Request body is empty. Send markdown text as the raw POST body.")
                return

            docx_bytes = convert_markdown_to_docx(markdown_text, mode=mode)
            docx_b64 = base64.b64encode(docx_bytes).decode('ascii')

            payload = json.dumps({
                "ok": True,
                "mode": mode,
                "ver": "v7-nokeep",
                "filename": "converted.docx",
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "size": len(docx_bytes),
                "data": docx_b64,
            })

            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('X-Powered-By', 'Python Cloud Function')
            self.end_headers()
            self.wfile.write(payload.encode('utf-8'))

        except Exception as exc:  # noqa: BLE001
            self._json_error(500, f"Conversion failed: {exc}")

    def do_GET(self):
        """Convenience info endpoint."""
        self.send_response(200)
        self.send_header('Content-Type', 'application/json')
        self.send_header('X-Powered-By', 'Python Cloud Function')
        self.end_headers()
        self.wfile.write(json.dumps({
            "route": "/api/convert/md-to-docx",
            "method": "POST",
            "description": "Send markdown text as the raw POST body; receive a JSON with base64-encoded .docx in 'data'.",
            "contentType": "text/plain (utf-8)",
            "modes": {
                "1": "basic — clean conversion with CJK font support",
                "2": "elegant — 1.5 line spacing, styled headings, code borders, zebra tables, hyperlinks, blockquote bars",
            },
            "usage": "POST /api/convert/md-to-docx?mode=2  (default mode=1)",
            "responseShape": {"ok": "bool", "mode": "int", "filename": "str", "mime": "str", "size": "int", "data": "base64 str"},
        }).encode('utf-8'))

    def _json_error(self, status, message):
        self.send_response(status)
        self.send_header('Content-Type', 'application/json')
        self.send_header('X-Powered-By', 'Python Cloud Function')
        self.end_headers()
        self.wfile.write(json.dumps({"ok": False, "error": message}).encode('utf-8'))
